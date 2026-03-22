import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
WORKDIR = os.getenv("WORKDIR", os.path.join(BASE_DIR, "workdir"))

def apply_gost_styles(doc: Document):
    section = doc.sections[0]
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.font.color.rgb = RGBColor(0, 0, 0)

    pfmt = style.paragraph_format
    pfmt.line_spacing = 1.5
    pfmt.space_before = Pt(0)
    pfmt.space_after = Pt(0)
    pfmt.first_line_indent = Cm(1.25)
    pfmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_cover_page(doc: Document, title_text: str = "Доклад по презентации"):
    p = doc.add_paragraph("МИНОБРНАУКИ РОССИИ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    p = doc.add_paragraph("\n\n\n")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(title_text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True

    doc.add_paragraph("\n\n\n")

    p = doc.add_paragraph(f"Автор: Студент")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p = doc.add_paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_page_break()


def add_black_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold = True

    if level <= 2:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_html_block(doc, html: str):
    soup = BeautifulSoup(html, "html.parser")

    for b in soup.find_all("b"):
        b.name = "strong"

    for elem in soup.children:
        if elem.name == "p":
            p = doc.add_paragraph()
            for child in elem.children:
                if child.name == "strong":
                    run = p.add_run(child.get_text())
                    run.bold = True
                else:
                    run = p.add_run(str(child))

                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif elem.name == "ul":
            for li in elem.find_all("li"):
                p = doc.add_paragraph(li.get_text(), style="List Bullet")
                for run in p.runs:
                    run.font.size = Pt(14)
                    run.font.color.rgb = RGBColor(0, 0, 0)


def extract_sources_from_slide_html(html: str) -> list[str]:
    soup = BeautifulSoup(html or "", "html.parser")
    sources = []

    p_sources = None
    for p in soup.find_all("p"):
        if "источники" in p.get_text(" ", strip=True).lower():
            p_sources = p
            break

    if not p_sources:
        return sources

    ul = p_sources.find_next("ul")
    if not ul:
        return sources

    for li in ul.find_all("li"):
        txt = li.get_text(" ", strip=True)
        if not txt:
            continue
        if "источники на слайде не указаны" in txt.lower():
            continue
        sources.append(txt)
    return sources


def build_docx_from_slides(json_path: str, output_path: str):
    with open(json_path, "r", encoding="utf-8") as f:
        slides = json.load(f)

    doc = Document()
    apply_gost_styles(doc)
    add_cover_page(doc, title_text="Доклад по презентации")

    all_sources: list[str] = []

    for slide in slides:
        slide_num = slide["slide"]
        html_text = slide["generated_html"]

        add_black_heading(doc, f"Слайд {slide_num}", level=1)
        add_html_block(doc, html_text)

        all_sources.extend(extract_sources_from_slide_html(html_text))

        doc.add_page_break()

    unique_sources = []
    seen = set()
    for s in all_sources:
        key = s.strip().lower()
        if key and key not in seen:
            seen.add(key)
            unique_sources.append(s.strip())

    add_black_heading(doc, "Список использованных источников", level=1)
    if unique_sources:
        for i, src in enumerate(unique_sources, start=1):
            p = doc.add_paragraph(f"{i}. {src}")
            for run in p.runs:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        add_paragraph(doc, "Источники не обнаружены.")

    doc.save(output_path)
    print("DOCX saved:", output_path)


if __name__ == "__main__":
    build_docx_from_slides(
        json_path=os.path.join(WORKDIR, "slides_report.json"),
        output_path=os.path.join(WORKDIR, "result.docx"),
    )
